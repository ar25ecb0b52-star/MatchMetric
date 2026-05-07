import os
import re 
import unicodedata
import logging

def extract_text(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = file_path.rsplit(".", 1)[-1].lower()

    raw = ""  

    if ext == "pdf":
        raw = _extract_from_pdf(file_path)
    elif ext == "docx":
        raw = _extract_from_docx(file_path)
    elif ext == "txt":
        raw = _extract_from_txt(file_path)
    else:
        raise ValueError(...)

    cleaned = clean_text(raw)

    if not cleaned:
        raise RuntimeError("No readable text found")
    
    return cleaned

def _extract_from_pdf(file_path:str)->str:
    import pdfplumber

    pages =[]
    with pdfplumber.open(file_path) as pdf:
        for i , page in enumerate(pdf.pages):
            page_text = page.extract_text(x_tolerance = 2, y_tolerance = 2)
            if page_text:
                pages.append(page_text)
            else:
                logging.debug("no text returned")

    return "\n\n".join(pages)

def _extract_from_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                parts.append("  |  ".join(row_cells))
    return "\n".join(parts)
    
def _extract_from_txt(file_path: str) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()
    
def clean_text(text: str) -> str:

    if not text:
        return ""

    
    text = unicodedata.normalize("NFKC", text)
    
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)

    text = re.sub(r"[•◦▪▸▹►‣⁃‐‑–—●○■□✓✗✘]", "-", text)

    
    lines = [line.strip() for line in text.splitlines()]
    cleaned_lines = []
    blank_count = 0
    for line in lines:
        if line == "":
            blank_count += 1
            if blank_count <= 1:
                cleaned_lines.append(line)
        else:
            blank_count = 0
            cleaned_lines.append(line)

    text = "\n".join(cleaned_lines)


    text = re.sub(r"[^\S\n]+", " ", text)

 
    return text.strip()


def get_word_count(text: str) -> int:
    """Return the number of words in the cleaned text."""
    return len(text.split())


def get_file_info(file_path: str) -> dict:
    """Return basic metadata about an uploaded file."""
    stat = os.stat(file_path)
    return {
        "filename": os.path.basename(file_path),
        "extension": file_path.rsplit(".", 1)[-1].lower(),
        "size_kb": round(stat.st_size / 1024, 2),
    }